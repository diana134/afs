"""The Schedule object used by the scheduling algorithm"""

import sys
# from random import randrange, shuffle
import datetime
import pickle
from docx import Document
from docx.shared import Inches, Pt
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.enum.style import WD_STYLE
# from docx.enum.style import WD_STYLE_TYPE

from databaseInteraction import dbInteractionInstance
# from scheduler import Scheduler

class Session(object):
    """Part of a Schedule"""
    def __init__(self, startDatetime=None, endDatetime=None):
        self.startDatetime = startDatetime
        self.endDatetime = endDatetime
        self.eventList = []

    def add(self, event):
        """Add an event to this session"""
        self.eventList.append(event)

    def remove(self, event):
        """Remove an event from this session"""
        self.eventList.remove(event)

    def filledTime(self):
        """Returns how much time is filled by the events"""
        filledTime = datetime.timedelta()
        for event in self.eventList:
            filledTime += event.totalTime
        return filledTime

    def emptyTime(self):
        """Returns how much empty time there is, will be negative if events go overtime"""
        totalTime = self.endDatetime - self.startDatetime
        emptyTime = totalTime - self.filledTime()
        return emptyTime

    def isFull(self):
        """Returns True if there is no more empty time"""
        if self.emptyTime().total_seconds() == 0:
            return True
        else:
            return False

    def hasTime(self, durationToFit):
        """Returns True if there is sufficient empty time to fit durationToFit"""
        if self.emptyTime() >= durationToFit:
            return True
        else:
            return False
            
    def export(self, csvFile):
        """Export this session to a csv.  The csvFile parameter must be a file with write permissions"""
        s = '"{startDate}","{endDate}","{numEvents} events"\n'.format(
            startDate=self.startDatetime.strftime('%Y/%m/%d %I:%M %p'),
            endDate=self.endDatetime.strftime('%Y/%m/%d %I:%M %p'),
            numEvents=len(self.eventList)
        )
        csvFile.write(s)
        for i in range(len(self.eventList)):
            e = self.eventList[i]
            csvFile.write('Event {n},'.format(n=i+1))
            e.export(csvFile)

    def toWordFile(self, document):
        """Create a docx for the printer, document parameter from docx module"""
        # Session start time (weekday, month day, year, hour:minute am/pm)
        s = "{0}".format(self.startDatetime.strftime('%A, %B %d, %Y %I:%M %p'))
        p = document.add_paragraph()
        r = p.add_run(s)
        r.bold = True
        r.caps = True
        for event in self.eventList:
            # event title (class number and name)
            eventTitle = "{0}    {1}".format(event.classNumber.replace(" ", ""), event.className)
            p = document.add_paragraph()
            p.add_run(eventTitle).bold = True
            event.toWordFile(document)

class Schedule(object):
    """Used by the scheduling algorithm"""
    def __init__(self, sessionDatetimes=None):
        self.sessions = []
        if sessionDatetimes is not None:
            for startDatetime, endDatetime in sessionDatetimes:
                self.sessions.append(Session(startDatetime, endDatetime))

    def findNextFit(self, durationToFit):
        """Returns the chronologically first session that isn't full, or None if they're all full"""
        # sort sessions ascending
        self.sessions.sort(key=lambda x: x.startDatetime)
        for session in self.sessions:
            print "trying to fit " + str(durationToFit) + " in session with emptyTime = " + str(session.emptyTime())
            if session.hasTime(durationToFit):
                return session
        return None

    def findWorstFit(self):
        """Returns the chronologically first session with the most empty time"""
        # Note: by its nature, this will put younger kids at the beginning of *each* day
        worstFitSession = None
        worstFitTime = datetime.timedelta()
        # sort sessions ascending
        self.sessions.sort(key=lambda x: x.startDatetime)
        for session in self.sessions:
            if session.emptyTime() > worstFitTime:
                worstFitSession = session
                worstFitTime = session.emptyTime()
        return worstFitSession

    def __str__(self):
        """Returns the string representation"""
        s = ""
        # sort sessions ascending by startDatetime
        self.sessions.sort(key=lambda x: x.startDatetime)
        for session in self.sessions:
            s = s + str(session.startDatetime) + '\n'
            for event in session.eventList:
                s = s + '\t' + event.classNumber + '\t' + event.className + '\n'
        return s

    def save(self, filename):
        """Save the schedule as a pickled blob and write it to the specified filename"""
        fout = open(filename, 'w')
        pickle.dump(self, fout)
        fout.close()
        
    def load(self, filename):
        """Load the schedule from a pickled blob"""
        fin = open(filename, 'r')
        loaded = pickle.load(fin)
        fin.close()
        
        # copy the data we loaded into self
        self.sessions = []
        for s in loaded.sessions:
            self.sessions.append(s)
            
    def export(self, filename):
        """Export the schedule as a reasonably-nicely formatted .csv file so they can play around with in in Excel"""
        fout = open(filename, 'w')
        for s in self.sessions:
            s.export(fout)
        fout.close()

    def toWordFile(self, filename):
        """Format schedule as a docx to send to the printer"""
        document = Document()
        p = document.add_paragraph()
        discipline = self.sessions[0].eventList[0].entries[0].discipline
        # Schedule details (DISCIPLINE - ##venue## \n Adjudicator - ##adjudicator##)
        # "##" fields denote things filled in by the user after the fact
        p.add_run(discipline.upper() + " - ##venue##\n").bold = True
        p.add_run("Adjudicator - ##adjudicator##\n").bold = True
        for s in self.sessions:
            s.toWordFile(document)
        document.save(filename)

    # def refresh(self):
    #     """Re-get all the entries from the db to refresh any changes"""
    #     entries = dbInteractionInstance.getAllEntriesInDiscipline(self.sessions[0].eventList[0].entries[0].discipline)
    #     newEventList = Scheduler.sortEntriesByClass(entries)

    #     for session in self.sessions:
    #         for event in session.eventList:
    #             for newEvent in newEventList:
    #                 # Find same events
    #                 if event.classNumber == newEvent.classNumber:
    #                     # TODO this destroys order inside events
    #                     # figure out how to preserve this (probably requires entries to store their db id's)
    #                     event.entries = newEvent.entries

    def printAdjudicationSheets(self, filename, adjudicator):
        """Creates a docx of all the adjudication sheets for the schedule"""
        print "PRINTING ADJUDICATION SHEETS"
        document = Document()

        for session in self.sessions:
            for event in session.eventList:
                classNumber = event.classNumber
                className = event.className.title()
                for entry in event.entries:
                    # Header
                    document.add_picture('2015header.jpg')
                    # header = document.add_paragraph()
                    # header.add_run().add_picture('logo.jpg', width=Inches(0.8))
                    # run = header.add_run('Rockwood Festival of the Arts\n')
                    # run.bold = True
                    # run.font.size = Pt(24)
                    # run = header.add_run('\t\t\t{location} - {year} Festival'.format(location=location, year=year))
                    # run.bold = True
                    # run.font.size = Pt(16)

                    # document.add_picture('logo.jpg', width=Inches(1.25))
                    # document.add_heading('Rockwood Festival of the Arts', 0)
                    # document.add_heading('{location} - {year} Festival'.format(location=location, year=year), level=1)
                    
                    # Entry information
                    participant = dbInteractionInstance.getParticipantFromId(entry.participantID)
                    pString = ""
                    try:
                        # Print soloist name
                        pString = "{0} {1}".format(participant.first, participant.last)
                    except Exception:
                        # Print list of participants in group
                        if len(participant.participants) > 0:
                            actualParticipants = []
                            tokens = participant.participants.split(',')
                            if tokens[0] != "":
                                for index in tokens:
                                    sp = dbInteractionInstance.getParticipantFromId(index)
                                    if sp.first != "":
                                        actualParticipants.append("{0} {1}".format(sp.first, sp.last))

                            # Correctly "comma-ify" the list of names
                            pString = ", ".join(actualParticipants)
                            index = pString.rfind(", ")
                            if index > -1:
                                pString = pString[:index] + " &" + pString[index+1:]
                        else:
                            # Print group name
                            pString = "{0}".format(participant.groupName)

                    selString = ""
                    for selection in entry.selections:
                        selString += "{0}".format(selection['title'])
                        if selection['titleOfMusical'] != "":
                            selString += " ({0})".format(selection['titleOfMusical'])
                        if selection['composerArranger'] != "":
                            selString += " - {0}".format(selection['composerArranger'])
                        selString += ", "
                    # Cut off last comma
                    index = selString.rfind(", ")
                    if index > -1:
                        selString = selString[:index]

                    info = document.add_paragraph()
                    # TODO fake some horizontal lines (maybe can do with custom style in ms word?)
                    run = info.add_run("\tClass: {0} {1}\n".format(classNumber, className))
                    run.bold = True
                    run.font.size = Pt(12)
                    run = info.add_run("\tParticipant(s): {0}\n".format(pString))
                    run.bold = True
                    run.font.size = Pt(16)
                    run = info.add_run("\tSelection(s): {0}\n".format(selString))
                    run.bold = True
                    run.font.size = Pt(16)

                    # Marking scale
                    marks = document.add_paragraph()
                    # marks.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = marks.add_run("STANDARD SCALE OF MARKS\n75 represents a performance which is good.\n80 represent Merit.\n85 represents Distinction. \n90 represents Honours.")
                    run.font.size = Pt(8)

                    markLine = document.add_paragraph()
                    # markLine.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    markLine.add_run("\nMark : ").font.size = Pt(8)
                    run = markLine.add_run(" "*30)
                    run.underline = True
                    run.font.size = Pt(8)

                    # Adjudicator
                    # footerStyle = document.styles["Footer"]
                    # adj = document.add_paragraph(style=footerStyle)
                    adj = document.add_paragraph()
                    # adj.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    adj.add_run("\n"*20)
                    run = adj.add_run(" "*100 + "\n")
                    run.underline = True
                    run.font.size = Pt(8)
                    run = adj.add_run("{adjudicator} - Adjudicator".format(adjudicator=adjudicator))
                    run.font.size = Pt(8)

                    document.add_page_break()

        document.save(filename)
